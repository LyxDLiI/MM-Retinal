from docx import Document
from collections import defaultdict
import re
import csv
import os

def match_book3(input_path, image_output_dir):
    document = Document(input_path)
    cn_result = defaultdict(str)
    en_result = defaultdict(str)

    pattern_cn_en = r'^(图([\d-]*).*)(Fig.*)'
    pattern_cn = r'^图([\d-]*).*'
    pattern_en = r'Fig[\. ]*([\d-]*).*'
    pattern_multi2one = r'[A-Z][\.．]'

    multipic_count = 0

    # Ensure the image output directory exists
    os.makedirs(image_output_dir, exist_ok=True)
    image_count = 0

    for para in document.paragraphs:
        match = re.search(pattern_cn_en, para.text)

        if match:
            id = match.group(2)
            cn_result[id] = match.group(1)
            en_result[id] = match.group(3)
            is_multipic_match = re.search(pattern_multi2one, match.group(1))
            if is_multipic_match:
                print(match.group(1))
            continue
        
        match = re.search(pattern_cn, para.text)
        if match:
            id = match.group(1)
            cn_result[id] = match.string
            is_multipic_match = re.search(pattern_multi2one, match.string)
            if is_multipic_match:
                print(match.string)
                multipic_count += 1
            continue
        
        match = re.search(pattern_en, para.text)
        if match:
            id = match.group(1)
            en_result[id] = match.string

    print("en_result", en_result)
    print("cn_result", cn_result)

    # Save images and get their names
    image_names = save_images(document, image_output_dir)

    return cn_result, en_result, multipic_count, image_names

def match_book4(input_path, image_output_dir):
    document = Document(input_path)
    en_result = defaultdict(str)

    pattern_en = r'Fig\.\s\d+\.\d+'
    pattern_multi2one = r'A[\.．]'

    multipic_count = 0

    # Ensure the image output directory exists
    os.makedirs(image_output_dir, exist_ok=True)

    for para in document.paragraphs:
        match = re.search(pattern_en, para.text)
        if match:
            id = match.group(0)
            en_result[id] = match.string

    print("en_result", en_result)

    # Save images and get their names
    image_names = save_images(document, image_output_dir)

    return en_result, multipic_count, image_names

def save_images(document, output_dir):
    import numpy as np
    from PIL import Image
    from io import BytesIO
    image_names = []
    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            image = rel.target_part.blob
            image_ext = rel.target_part.content_type.split('/')[-1]
            image_filename = f"figure_{len(image_names) + 1}.{image_ext}"
            image_path = os.path.join(output_dir, image_filename)
            
            # Convert the image blob to a numpy array to check unique colors
            img = Image.open(BytesIO(image))
            img_array = np.array(img)

            # Check if the image has a significant number of unique colors
            unique_colors = np.unique(img_array.reshape(-1, img_array.shape[2]), axis=0)
            if len(unique_colors) > 30:  # Assuming a threshold of 10 unique colors
                with open(image_path, 'wb') as f:
                    f.write(image)
                image_names.append(image_filename)
                print(f"Saved image as {image_filename}")
            else:
                print(f"Skipped saving image {image_filename} due to low unique colors")

    return image_names

def write_into_csv_book4(output_path, match_info):
    en_result, multipic_count, image_names = match_info
    pattern_multi2one = r'[A-Z][\.．]'
    pattern = r'[A-Z]'

    with open(output_path, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Image_ID", "cn_caption", "en_caption", "is_multipic"])

        for i, key in enumerate(en_result.keys()):
            m_match = re.search(pattern_multi2one, en_result[key])
            is_multipic = ''
            if m_match:
                is_multipic = 'Y'
            image_id = image_names[i] if i < len(image_names) else f'figure_{i + 1}'
            writer.writerow([image_id, '', en_result[key], is_multipic])

def write_into_csv_book3(output_path, match_info):
    cn_result, en_result, multipic_count, image_names = match_info
    pattern_multi2one = r'[A-Z][\.．]'

    with open(output_path, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Image_ID", "cn_caption", "en_caption", "is_multipic"])

        for i, key in enumerate(cn_result.keys()):
            m_match = re.search(pattern_multi2one, cn_result[key])
            is_multipic = ''
            if m_match:
                is_multipic = 'Y'
            image_id = image_names[i] if i < len(image_names) else f'figure_{i + 1}'
            writer.writerow([image_id, cn_result[key], en_result[key], is_multipic])

if __name__ == "__main__":
    book_path = '../book1/'  # replace
    total_part_number = 1    # replace (the number of sub-PDFs)
    for i in range(total_part_number):
        input_docx_path = book_path + 'texts/part' + str(i + 1) + '.docx'
        output_csv_path = book_path + 'caption/part' + str(i + 1) + '.csv'
        image_output_dir = book_path + 'images/part' + str(i + 1) + '/'
        match_result = match_book4(input_docx_path, image_output_dir)  # replace
        write_into_csv_book4(output_csv_path, match_result)            # replace
